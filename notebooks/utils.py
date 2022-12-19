from os import mkdir
from os.path import dirname, exists, abspath

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def profile_to_word_doc(profile_id: int, df_profiles: pd.DataFrame, df_skills: pd.DataFrame,
                        df_experiences: pd.DataFrame, out_file_name: str = "profile.docx"):
    """
    Convert a profile ot a Word document
    :param profile_id: id of profile to convert
    :param df_profiles: dataframe of profiles
    :param df_skills: dataframe of profile skills
    :param df_experiences: dataframe of profile experiences
    :param out_file_name: generated document path/name (needs to end in .docx)
    """
    jobtitle = df_profiles.loc[df_profiles["id_profile"] == profile_id, "jobtitle"].item()
    skills = df_skills.loc[df_skills["id_profile"] == profile_id, "skill"].to_list()
    experiences = df_experiences.loc[df_experiences["id_profile"] == profile_id, ["title", "description"]].T.to_dict()

    document = Document()
    document.add_heading(f"CV pour: {jobtitle}", 0)
    document.add_heading("Experiences:", 1)
    for experience in experiences.values():
        document.add_paragraph(experience["title"], style='ListBullet')
        if not isinstance(experience['description'], str):
            continue
        paragraph = document.add_paragraph(experience['description'])
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.add_heading("Competences:", 1)
    for skill in skills:
        document.add_paragraph(skill, style='ListBullet')

    # Save
    path = abspath(out_file_name)
    if not exists(dirname(path)):
        mkdir(dirname(path))
    document.save(path)


def offer_to_word_doc(jobtitle: str, description: str, contract: str, company: str, location: str,
                      out_file_name: str = "offre.docx"):
    """
    Convert an offer to a Word document
    :param jobtitle: offer job title
    :param description: offer description
    :param contract: offer contract type
    :param company:  offer company
    :param location: company location
    :param out_file_name: generated document path/name (needs to end in .docx)
    """
    document = Document()
    document.add_heading(f"Offre pour: {jobtitle}", 0)
    if isinstance(company, str):
        document.add_heading("Entreprise:", 1)
        document.add_paragraph(company)
    if isinstance(contract, str):
        document.add_heading("Type de contrat:", 1)
        document.add_paragraph(contract)
    if isinstance(location, str):
        document.add_heading("Lieu de travail:", 1)
        document.add_paragraph(location)
    document.add_heading("Description de l'offre:", 1)
    for desc_part in description.split("\n"):
        if len(desc_part) == 0:
            continue
        paragraph = document.add_paragraph(desc_part)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save
    path = abspath(out_file_name)
    if not exists(dirname(path)):
        mkdir(dirname(path))
    document.save(path)


if __name__ == "__main__":
    def example():
        # Example to convert an offer to a Word document
        df_offers = pd.read_csv('offers.csv')
        offer_to_word_doc(**df_offers.iloc[1, :].to_dict(), out_file_name="offre.docx")

        # Example to convert a profile to a Word document
        df_profiles = pd.read_csv('employee_profiles.csv')
        df_skills = pd.read_csv('employee_skills.csv')
        df_experiences = pd.read_csv('employee_experiences.csv')
        profile_id = 49282
        profile_to_word_doc(profile_id, df_profiles, df_skills, df_experiences)


    example()
